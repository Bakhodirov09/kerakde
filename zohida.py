import asyncio
from aiogram import Bot, Dispatcher, types
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from aiogram.dispatcher import FSMContext
from aiogram.dispatcher.filters.state import State, StatesGroup
from aiogram.utils import executor
from docx import Document

TOKEN = "8045362943:AAGhvDXgKEfXp1N15oeI_XorUMQJp_lVCHw"

bot = Bot(token=TOKEN)
dp = Dispatcher(bot, storage=MemoryStorage())

class ProductForm(StatesGroup):
    name = State()
    amount = State()

@dp.message_handler(commands=["start"])
async def start(message: types.Message):
    await message.answer("Mahsulot nomini kiriting:")
    await ProductForm.name.set()

@dp.message_handler(state=ProductForm.name)
async def get_name(message: types.Message, state: FSMContext):
    await state.update_data(name=message.text)
    await message.answer("Miqdorini kiriting:")
    await ProductForm.amount.set()

@dp.message_handler(state=ProductForm.amount)
async def get_amount(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data["name"]
    amount = message.text

    # Faylni yaratish yoki ochish
    file_name = "products.docx"
    try:
        doc = Document(file_name)
    except:
        doc = Document()
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "T/r"
        hdr_cells[1].text = "Mahsulot nomi"
        hdr_cells[2].text = "Miqdori"
        doc.save(file_name)
        doc = Document(file_name)

    table = doc.tables[0]
    row = table.add_row().cells
    row[0].text = str(len(table.rows) - 1)
    row[1].text = name
    row[2].text = amount
    doc.save(file_name)

    await message.answer("✅ Ma'lumot saqlandi!")
    await state.finish()

@dp.message_handler(commands=["file"])
async def send_file(message: types.Message):
    if os.path.exists(FILE_NAME):
        await message.answer_document(open('products.docx', "rb"))
    else:
        await message.answer("❌ Hozircha fayl mavjud emas!")

if __name__ == "__main__":
    executor.start_polling(dp, skip_updates=True)
